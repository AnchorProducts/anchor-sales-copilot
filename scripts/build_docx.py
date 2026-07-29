#!/usr/bin/env python3
"""Render SOP.md (+ SITEMAP.md) into a Word .docx.

    python3 scripts/build_docx.py

Reads the same two markdown files the HTML build uses, and the figure manifest
it emits (docs/shots.json), so the Word and web versions never drift.

Images are the reason this isn't a one-liner. The raw screenshots are full
retina captures — several are 500KB and some are tall phone portraits that would
run past a full page if placed at column width. Every image is therefore:
  * downscaled to MAX_PX on its long edge into a build cache (originals in
    docs/images/ are never touched), and
  * fitted inside a MAX_W x MAX_H box, so portrait shots shrink to fit rather
    than blowing out the page.

Requires python-docx and macOS `sips` (for the downscale; skipped if absent).
"""

import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip3 install python-docx")

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
IMAGES = DOCS / "images"
CACHE = DOCS / ".docx-images"          # downscaled copies; safe to delete
OUT = DOCS / "Anchor-Sales-Copilot-SOP.docx"

MAX_PX = 1000        # long-edge pixels after downscale
JPEG_QUALITY = 75    # see prepared() — PNG re-encoding made the cache bigger
MAX_W = Inches(5.6)  # never wider than this
MAX_H = Inches(3.6)  # never taller — keeps tall phone shots on the page

GREEN = RGBColor(0x0B, 0x5D, 0x2E)
GREY = RGBColor(0x5D, 0x6B, 0x64)
MONO = "Consolas"


# ── images ──────────────────────────────────────────────────────────────────
def png_size(path: Path):
    """(width, height) straight from the PNG/JPEG header — no Pillow needed."""
    data = path.read_bytes()
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return int.from_bytes(data[16:20], "big"), int.from_bytes(data[20:24], "big")
    if data[:2] == b"\xff\xd8":  # JPEG
        i = 2
        while i < len(data) - 9:
            if data[i] != 0xFF:
                i += 1
                continue
            marker, seglen = data[i + 1], int.from_bytes(data[i + 2 : i + 4], "big")
            if marker in (0xC0, 0xC1, 0xC2, 0xC3):
                return (
                    int.from_bytes(data[i + 7 : i + 9], "big"),
                    int.from_bytes(data[i + 5 : i + 7], "big"),
                )
            i += 2 + seglen
    return None


def prepared(name: str):
    """Downscaled JPEG copy of docs/images/<name>, or None when absent.

    JPEG, not PNG: these are mostly narrow phone screenshots already close to
    the target pixel size, so resampling a PNG barely shrinks the dimensions
    while re-encoding it *worse* — the cache came out bigger than the originals
    (8.9MB from 7.3MB). Re-encoding to JPEG q75 at the same dimensions gives
    2.8MB, which is the difference between a Word file you can email and one you
    can't. At 5.6in wide the artifacts aren't visible.
    """
    src = IMAGES / name
    if not src.exists():
        return None
    CACHE.mkdir(exist_ok=True)
    dst = CACHE / (Path(name).stem + ".jpg")
    if not dst.exists() or dst.stat().st_mtime < src.stat().st_mtime:
        shutil.copy2(src, dst)
        if shutil.which("sips"):
            subprocess.run(
                ["sips", "-s", "format", "jpeg", "-s", "formatOptions", str(JPEG_QUALITY),
                 "--resampleHeightWidthMax", str(MAX_PX), str(dst)],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=False,
            )
    return dst


def fitted(path: Path):
    """Width/height that fits the image inside MAX_W x MAX_H, preserving aspect."""
    dims = png_size(path)
    if not dims:
        return MAX_W, None
    w, h = dims
    scale = min(MAX_W / Inches(1) / (w / 96), MAX_H / Inches(1) / (h / 96))
    return Inches((w / 96) * scale), Inches((h / 96) * scale)


# ── inline markdown → runs ──────────────────────────────────────────────────
TOKEN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|(?<![*\w])\*[^*\n]+\*|\[[^\]]+\]\([^)\s]+\))")


def add_runs(par, text: str, bold=False, italic=False, size=None):
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = MONO
            r.font.size = Pt((size or 10.5) - 1)
        elif part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = par.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)\s]+)\)", part)
            r = par.add_run(m.group(1) if m else part)
            r.font.color.rgb = GREEN
            r.underline = True
        else:
            r = par.add_run(part)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if size:
            r.font.size = Pt(size)


# ── block parsing (same subset the HTML build supports) ─────────────────────
def blocks(md: str):
    out, buf = [], []
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("```"):
            if buf:
                out.append("\n".join(buf)); buf = []
            fence = [ln]
            i += 1
            while i < len(lines):
                fence.append(lines[i])
                if lines[i].startswith("```"):
                    break
                i += 1
            out.append("\n".join(fence))
        elif not ln.strip():
            if buf:
                out.append("\n".join(buf)); buf = []
        elif re.match(r"^#{1,6}\s", ln) or re.match(r"^(---|\*\*\*)\s*$", ln) or ln.strip() == "\\newpage":
            if buf:
                out.append("\n".join(buf)); buf = []
            out.append(ln.strip())
        else:
            buf.append(ln)
        i += 1
    if buf:
        out.append("\n".join(buf))
    return out


def add_table(doc, rows):
    def cells(r):
        return [c.strip() for c in r.strip().strip("|").split("|")]

    head = cells(rows[0])
    body = [cells(r) for r in rows[2:]]
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(head):
        cell = t.rows[0].cells[i]
        cell.text = ""
        add_runs(cell.paragraphs[0], h, bold=True, size=9.5)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "EFF4F1")
        cell._tc.get_or_add_tcPr().append(shade)
    for row in body:
        cs = t.add_row().cells
        for i in range(len(head)):
            cs[i].text = ""
            add_runs(cs[i].paragraphs[0], row[i] if i < len(row) else "", size=9.5)
    doc.add_paragraph()


def add_figure(doc, shot, n):
    """A screenshot, or a visible placeholder when the PNG hasn't been taken."""
    img = prepared(shot["file"])
    if img:
        w, h = fitted(img)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img), width=w, height=h)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(cap, f"Fig. {n} — {shot['caption']}", size=9)
        for r in cap.runs:
            r.font.color.rgb = GREY
    else:
        p = doc.add_paragraph()
        add_runs(p, f"[ Fig. {n} — screenshot needed: {shot['file']} ]", italic=True, size=9)
        for r in p.runs:
            r.font.color.rgb = GREY
        d = doc.add_paragraph()
        add_runs(d, f"Where: {shot['where']} — {shot['what']}", size=9)
        for r in d.runs:
            r.font.color.rgb = GREY
    doc.add_paragraph()


def render(doc, md, shots, counter, headings_are_sitemap=False):
    used = counter["used"]
    for raw in blocks(md):
        lines = raw.split("\n")

        heading = re.match(r"^(#{1,6})\s+(.*)$", raw)
        is_table = (
            lines[0].lstrip().startswith("|")
            and len(lines) > 1
            and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[1])
        )

        # One if/elif chain, so EVERY block falls through to the figure loop at
        # the bottom. (Early `continue`s here previously skipped it, which lost
        # every figure anchored to a heading or a table — most of them.)
        if raw in ("---", "***"):
            pass
        elif raw == "\\newpage":
            doc.add_page_break()
        elif raw.startswith("```"):
            body = "\n".join(lines[1:-1])
            p = doc.add_paragraph()
            r = p.add_run(body)
            r.font.name = MONO
            r.font.size = Pt(9)
        elif heading:
            level, text = len(heading.group(1)), heading.group(2).strip()
            if level == 1 and not headings_are_sitemap and counter["seen_h1"]:
                doc.add_page_break()
            if level == 1:
                counter["seen_h1"] = True
            if level == 2 and not headings_are_sitemap and text.startswith("Page "):
                doc.add_page_break()
            h = doc.add_heading(level=min(level + (1 if headings_are_sitemap else 0), 4))
            add_runs(h, re.sub(r"[*`]", "", text))
        elif is_table:
            add_table(doc, [l for l in lines if l.strip().startswith("|")])
        elif all(re.match(r"^\s*>", l) for l in lines):
            inner = " ".join(l.lstrip("> ").rstrip() for l in lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_runs(p, inner, size=10)
            for r in p.runs:
                r.font.color.rgb = GREEN
        elif all(re.match(r"^\s*-\s\[[ xX]\]\s", l) for l in lines):
            for l in lines:
                mm = re.match(r"^\s*-\s\[([ xX])\]\s(.*)$", l)
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, ("☑ " if mm.group(1).lower() == "x" else "☐ ") + mm.group(2))
        elif re.match(r"^\s*[-*]\s", lines[0]):
            items = []
            for l in lines:
                if re.match(r"^\s*[-*]\s", l):
                    items.append(re.sub(r"^\s*[-*]\s", "", l))
                elif items:
                    items[-1] += " " + l.strip()
            for it in items:
                add_runs(doc.add_paragraph(style="List Bullet"), it)
        elif re.match(r"^\s*\d+\.\s", lines[0]):
            items = []
            for l in lines:
                if re.match(r"^\s*\d+\.\s", l):
                    items.append(re.sub(r"^\s*\d+\.\s", "", l))
                elif items:
                    items[-1] += " " + l.strip()
            for it in items:
                add_runs(doc.add_paragraph(style="List Number"), it)
        elif raw.strip() == "<!--SITEMAP-->":
            counter["sitemap_here"] = True
        elif raw.lstrip().startswith("<"):
            continue
        else:
            add_runs(doc.add_paragraph(), " ".join(l.strip() for l in lines))

        # Figures attach to the block whose text anchors them.
        for shot in shots:
            if shot["file"] in used:
                continue
            if shot["after"] in raw:
                used.add(shot["file"])
                counter["n"] += 1
                add_figure(doc, shot, counter["n"])


def main():
    sop = (ROOT / "SOP.md").read_text()
    sitemap = (ROOT / "SITEMAP.md").read_text()
    shots = json.loads((DOCS / "shots.json").read_text())

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    counter = {"n": 0, "used": set(), "seen_h1": False, "sitemap_here": False}

    # SOP up to the sitemap marker, then the sitemap, then anything after.
    head, _, tail = sop.partition("<!--SITEMAP-->")
    render(doc, head, shots, counter)
    render(doc, sitemap, [], counter, headings_are_sitemap=True)
    if tail.strip():
        render(doc, tail, shots, counter)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)

    have = sum(1 for s in shots if (IMAGES / s["file"]).exists())
    size = OUT.stat().st_size / 1_048_576
    print(f"✓ {OUT.relative_to(ROOT)}")
    print(f"  {counter['n']} figures placed ({have} images available) · {size:.1f} MB")


if __name__ == "__main__":
    main()
